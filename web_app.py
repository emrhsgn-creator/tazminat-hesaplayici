import streamlit as st
from datetime import datetime, timedelta
import pandas as pd
from docx import Document
from io import BytesIO

# --- 1. MASAÜSTÜNDEKİ BİREBİR VERİTABANI ---
TRH_2010_ERKEK = {
    0: 71.93, 1: 72.35, 2: 71.42, 3: 70.47, 4: 69.52, 5: 68.57, 6: 67.60, 7: 66.63, 8: 65.66, 9: 64.68,
    10: 63.70, 11: 62.72, 12: 61.74, 13: 60.76, 14: 59.78, 15: 58.80, 16: 57.84, 17: 56.87, 18: 55.91, 19: 54.95,
    20: 53.99, 21: 53.04, 22: 52.09, 23: 51.14, 24: 50.19, 25: 49.24, 26: 48.28, 27: 47.33, 28: 46.37, 29: 45.41,
    30: 44.45, 31: 43.50, 32: 42.54, 33: 41.58, 34: 40.62, 35: 39.67, 36: 38.72, 37: 37.77, 38: 36.81, 39: 35.87,
    40: 34.93, 41: 33.99, 42: 33.05, 43: 32.12, 44: 31.19, 45: 30.27, 46: 29.36, 47: 28.46, 48: 27.56, 49: 26.67,
    50: 25.79, 51: 24.93, 52: 24.08, 53: 23.24, 54: 22.42, 55: 21.61, 56: 20.82, 57: 20.04, 58: 19.28, 59: 18.54,
    60: 17.81, 61: 17.10, 62: 16.41, 63: 15.74, 64: 15.08, 65: 14.44, 66: 13.82, 67: 13.22, 68: 12.64, 69: 12.08,
    70: 11.54, 71: 11.01, 72: 10.50, 73: 10.01, 74: 9.54, 75: 9.08, 76: 8.64, 77: 8.21, 78: 7.80, 79: 7.41,
    80: 7.03, 81: 6.66, 82: 6.31, 83: 5.98, 84: 5.66, 85: 5.35, 86: 5.06, 87: 4.78, 88: 4.51, 89: 4.25,
    90: 4.00, 91: 3.76, 92: 3.53, 93: 3.31, 94: 3.10, 95: 2.90, 96: 2.71, 97: 2.53, 98: 2.36, 99: 2.20, 100: 2.05
}

TRH_2010_KADIN = {
    0: 78.02, 1: 77.66, 2: 76.68, 3: 75.70, 4: 74.72, 5: 73.73, 6: 72.74, 7: 71.75, 8: 70.76, 9: 69.76,
    10: 68.77, 11: 67.78, 12: 66.78, 13: 65.79, 14: 64.79, 15: 63.80, 16: 62.81, 17: 61.82, 18: 60.83, 19: 59.84,
    20: 58.85, 21: 57.86, 22: 56.88, 23: 55.89, 24: 54.90, 25: 53.92, 26: 52.93, 27: 51.95, 28: 50.97, 29: 49.98,
    30: 49.00, 31: 48.02, 32: 47.04, 33: 46.06, 34: 45.08, 35: 44.10, 36: 43.12, 37: 42.15, 38: 41.17, 39: 40.20,
    40: 39.23, 41: 38.26, 42: 37.30, 43: 36.34, 44: 35.38, 45: 34.43, 46: 33.48, 47: 32.54, 48: 31.60, 49: 30.67,
    50: 29.74, 51: 28.82, 52: 27.91, 53: 27.01, 54: 26.12, 55: 25.24, 56: 24.38, 57: 23.53, 58: 22.69, 59: 21.86,
    60: 21.04, 61: 20.24, 62: 19.45, 63: 18.67, 64: 17.91, 65: 17.16, 66: 16.42, 67: 15.70, 68: 14.99, 69: 14.30,
    70: 13.62, 71: 12.96, 72: 12.31, 73: 11.68, 74: 11.06, 75: 10.46, 76: 9.88, 77: 9.31, 78: 8.76, 79: 8.23,
    80: 7.71, 81: 7.21, 82: 6.73, 83: 6.27, 84: 5.82, 85: 5.39, 86: 4.98, 87: 4.59, 88: 4.22, 89: 3.86,
    90: 3.52, 91: 3.20, 92: 2.90, 93: 2.62, 94: 2.36, 95: 2.12, 96: 1.90, 97: 1.70, 98: 1.52, 99: 1.36, 100: 1.22
}

YILLIK_NET_ASGARI_UCRETLER = {
    2015: 1000.54, 2016: 1300.99, 2017: 1404.06, 2018: 1603.12, 2019: 2020.90, 2020: 2324.71,
    2021: 2825.90, 2022: 4876.50, 2023: 9954.50, 2024: 17002.12, 2025: 22000.00, 2026: 28075.50
}

def create_word_doc(rapor_icerik):
    doc = Document()
    doc.add_heading('AKTÜERYA BİLİRKİŞİ RAPORU', 0)
    for line in rapor_icerik.split('\n'):
        doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 2. MASAÜSTÜ (EXE) HESAPLAMA MOTORUNUN HARFİYEN KOPYASI ---
class AktueryaHesaplayici:
    def __init__(self, isim, cinsiyet, dogum_tarihi, kaza_tarihi, ise_baslama_tarihi, hesaplama_tarihi, 
                 maluliyet_orani, kusursuzluk_orani, maas_gecmisi, guncel_asgari):
        
        self.isim = isim
        self.cinsiyet = cinsiyet
        # EXE mantığındaki saniye kırılımlı hesaplamalar bozulmasın diye tekrar strptime kullanıyoruz
        self.dogum_tarihi = datetime.strptime(dogum_tarihi, "%d.%m.%Y")
        self.kaza_tarihi = datetime.strptime(kaza_tarihi, "%d.%m.%Y")
        self.ise_baslama_tarihi = datetime.strptime(ise_baslama_tarihi, "%d.%m.%Y")
        self.hesaplama_tarihi = datetime.strptime(hesaplama_tarihi, "%d.%m.%Y")
        
        self.kaza_yasi = self.kaza_tarihi.year - self.dogum_tarihi.year - ((self.kaza_tarihi.month, self.kaza_tarihi.day) < (self.dogum_tarihi.month, self.dogum_tarihi.day))
        if self.kaza_yasi < 0: self.kaza_yasi = 0
        if self.kaza_yasi > 100: self.kaza_yasi = 100
            
        if self.cinsiyet == "Erkek":
            self.bakiye_omur = TRH_2010_ERKEK[self.kaza_yasi]
        else:
            self.bakiye_omur = TRH_2010_KADIN[self.kaza_yasi]
        
        self.emeklilik_tarihi = self._yil_ekle(self.dogum_tarihi, 60)
        self.tahmini_olum_tarihi = self.kaza_tarihi + timedelta(days=self.bakiye_omur * 365.25)
        
        self.maluliyet_orani = maluliyet_orani / 100.0
        self.kusursuzluk_orani = kusursuzluk_orani
        self.maas_gecmisi = maas_gecmisi
        self.guncel_asgari = guncel_asgari
        
        self.rapor_metni = ""

    def _yil_ekle(self, tarih, yil):
        try:
            return tarih.replace(year=tarih.year + yil)
        except ValueError:
            return tarih + (datetime(tarih.year + yil, 1, 1) - datetime(tarih.year, 1, 1))

    def gunluk_ucret(self, aylik_ucret):
        return aylik_ucret / 30.0

    def hesapla(self):
        self.rapor_metni = f"MADDİ TAZMİNAT (İŞ GÖREMEZLİK) AKTÜERYA HESAPLAMA RAPORU\n"
        self.rapor_metni += f"{'='*70}\n"
        self.rapor_metni += f"Mağdur: {self.isim} ({self.cinsiyet})\n"
        self.rapor_metni += f"Doğum Tarihi: {self.dogum_tarihi.strftime('%d.%m.%Y')} | Kaza Tarihi: {self.kaza_tarihi.strftime('%d.%m.%Y')}\n"
        self.rapor_metni += f"Kaza Tarihindeki Yaşı: {self.kaza_yasi} Yaş\n"
        self.rapor_metni += f"TRH-2010 Bakiye Ömür: {self.bakiye_omur} Yıl\n"
        self.rapor_metni += f"Emeklilik Tarihi (60 Yaş): {self.emeklilik_tarihi.strftime('%d.%m.%Y')}\n"
        self.rapor_metni += f"Tahmini Vefat Tarihi: {self.tahmini_olum_tarihi.strftime('%d.%m.%Y')}\n"
        self.rapor_metni += f"Maluliyet: %{self.maluliyet_orani*100} | Kusursuzluk Çarpanı: {self.kusursuzluk_orani}\n"
        self.rapor_metni += f"{'='*70}\n\n"

        genel_toplam = 0

        self.rapor_metni += "1. GEÇİCİ İŞ GÖREMEZLİK DÖNEMİ (Maluliyet: %100)\n"
        self.rapor_metni += "-"*70 + "\n"
        gecici_toplam = self._donem_hesapla(self.kaza_tarihi, self.ise_baslama_tarihi, 1.0, self.kusursuzluk_orani, progresif=False)
        genel_toplam += gecici_toplam
        self.rapor_metni += f"Geçici İş Göremezlik Toplamı: {gecici_toplam:,.2f} TL\n\n"

        self.rapor_metni += "2. SÜREKLİ İŞ GÖREMEZLİK - BİLİNEN DÖNEM\n"
        self.rapor_metni += "-"*70 + "\n"
        bilinen_toplam = self._donem_hesapla(self.ise_baslama_tarihi, self.hesaplama_tarihi, self.maluliyet_orani, self.kusursuzluk_orani, progresif=False)
        genel_toplam += bilinen_toplam
        self.rapor_metni += f"Bilinen Dönem Toplamı: {bilinen_toplam:,.2f} TL\n\n"

        self.rapor_metni += "3. SÜREKLİ İŞ GÖREMEZLİK - BİLİNMEYEN DÖNEM (Progresif Rant: %10 Artırım / %10 İskonto)\n"
        self.rapor_metni += "-"*70 + "\n"
        
        if self.maas_gecmisi:
            taban_maas = self.maas_gecmisi[max(self.maas_gecmisi.keys())]
        else:
            taban_maas = self.guncel_asgari

        bilinmeyen_toplam = self._donem_hesapla(self.hesaplama_tarihi, self.emeklilik_tarihi, self.maluliyet_orani, self.kusursuzluk_orani, progresif=True, taban_maas=taban_maas)
        genel_toplam += bilinmeyen_toplam
        self.rapor_metni += f"Bilinmeyen Dönem Toplamı: {bilinmeyen_toplam:,.2f} TL\n\n"

        self.rapor_metni += "4. PASİF DÖNEM (Progresif Rant - Asgari Ücret Üzerinden)\n"
        self.rapor_metni += "-"*70 + "\n"
        pasif_toplam = self._donem_hesapla(self.emeklilik_tarihi, self.tahmini_olum_tarihi, self.maluliyet_orani, self.kusursuzluk_orani, progresif=True, taban_maas=self.guncel_asgari, baslangic_n=max(1, self.emeklilik_tarihi.year - self.hesaplama_tarihi.year + 1))
        genel_toplam += pasif_toplam
        self.rapor_metni += f"Pasif Dönem Toplamı: {pasif_toplam:,.2f} TL\n\n"

        self.rapor_metni += f"{'='*70}\n"
        self.rapor_metni += f"GENEL TOPLAM TAZMİNAT: {genel_toplam:,.2f} TL\n"
        self.rapor_metni += f"{'='*70}\n\n"
        
        kapanis = ("İşbu aktüerya hesaplama raporu, dosya muhteviyatındaki verilere, güncel Yargıtay içtihatlarına, "
                   "TRH-2010 yaşam tablosuna ve progresif rant (peşin değer) ilkelerine uygun olarak tarafımca tanzim edilmiştir.\n"
                   "Takdiri Sayın Mahkemeye ait olmak üzere saygılarımla arz ederim.")
        self.rapor_metni += kapanis

        return self.rapor_metni

    def _donem_hesapla(self, baslangic, bitis, maluliyet, kusur, progresif=False, taban_maas=None, baslangic_n=1):
        if baslangic >= bitis: return 0
            
        toplam_tutar = 0
        gecerli_tarih = baslangic
        yil_sayaci = baslangic_n

        if progresif:
            self.rapor_metni += f"{'YIL':<6} | {'GÜN':<4} | {'ARTIRIMLI MAAŞ':<16} | {'İSKONTO KATS.':<13} | {'MALUL.':<6} | {'PEŞİN DEĞER (TL)':<15}\n"
            self.rapor_metni += "-"*80 + "\n"

        while gecerli_tarih < bitis:
            yil = gecerli_tarih.year
            yil_sonu = datetime(yil, 12, 31)
            donem_bitisi = min(yil_sonu, bitis - timedelta(days=1))
            gun_sayisi = (donem_bitisi - gecerli_tarih).days + 1

            if progresif:
                artirim_katsayisi = (1.10) ** yil_sayaci
                iskonto_katsayisi = 1 / ((1.10) ** yil_sayaci)
                artirimli_maas = taban_maas * artirim_katsayisi
                gunluk_ucret = self.gunluk_ucret(artirimli_maas)
                pesin_deger = gun_sayisi * gunluk_ucret * iskonto_katsayisi * maluliyet * kusur
                
                self.rapor_metni += f"{yil:<6} | {gun_sayisi:<4} | {artirimli_maas:,.2f} TL    | {iskonto_katsayisi:.4f}        | %{maluliyet*100:<4} | {pesin_deger:,.2f} TL\n"
                toplam_tutar += pesin_deger
                yil_sayaci += 1
            else:
                if yil in self.maas_gecmisi:
                    maas = self.maas_gecmisi[yil]
                else:
                    maas = YILLIK_NET_ASGARI_UCRETLER.get(yil, self.guncel_asgari)
                    
                tutar = gun_sayisi * self.gunluk_ucret(maas) * maluliyet * kusur
                self.rapor_metni += f"{yil} Yılı: {gun_sayisi} Gün x {maas:,.2f} TL x %{maluliyet*100} = {tutar:,.2f} TL\n"
                toplam_tutar += tutar

            gecerli_tarih = donem_bitisi + timedelta(days=1)

        return toplam_tutar

# --- 3. STREAMLIT WEB ARAYÜZÜ ---
st.set_page_config(page_title="Aktüerya Bilirkişisi", page_icon="⚖️", layout="wide")
st.title("⚖️ Aktüerya Bilirkişi Paneli")

c1, c2 = st.columns(2)

with c1:
    isim = st.text_input("İsim Soyisim:", "Ahmet Mehmet")
    cinsiyet = st.selectbox("Cinsiyet:", ["Erkek", "Kadın"])
    dogum = st.date_input("Doğum Tarihi:", value=datetime(1996, 12, 20), min_value=datetime(1900, 1, 1), max_value=datetime.today())
    kaza = st.date_input("Kaza Tarihi:", value=datetime(2023, 11, 4), min_value=datetime(1900, 1, 1), max_value=datetime(2100, 1, 1))
    
with c2:
    ise_baslama = st.date_input("İşe Başlama Tarihi:", value=datetime(2024, 5, 4), min_value=datetime(1900, 1, 1), max_value=datetime(2100, 1, 1))
    hesap = st.date_input("Hesaplama Tarihi (Bugün):", value=datetime.today(), min_value=datetime(1900, 1, 1), max_value=datetime(2100, 1, 1))
    mal = st.number_input("Maluliyet Oranı (%):", min_value=0.0, max_value=100.0, value=5.0, step=1.0)
    kusur = st.number_input("Kusursuzluk (Haklılık) Oranı (Örn: 1.0 veya 0.75):", min_value=0.0, max_value=1.0, value=1.0, step=0.01, format="%.3f")
    asgari = st.number_input("Güncel Asgari Ücret (TL):", min_value=0.0, value=28075.50, step=100.0)

st.subheader("📊 Maaş Geçmişi")
st.info("💡 Asgari ücretle çalışanlar için tabloyu bomboş bırakın. Farklı bir maaş varsa alttaki tabloya satır ekleyebilirsiniz.")

yillar = list(range(2010, 2035))
df_in = pd.DataFrame(columns=["Yıl", "Maaş (TL)"]) 
edit_df = st.data_editor(df_in, num_rows="dynamic", use_container_width=True, 
                         column_config={"Yıl": st.column_config.SelectboxColumn("Yıl", options=yillar), 
                                        "Maaş (TL)": st.column_config.NumberColumn("Maaş (TL)", format="%.2f")})

if st.button("Bilirkişi Raporunu Oluştur", type="primary", use_container_width=True):
    # Webden gelen tarih verilerini, tam da EXE programının istediği (04.11.2023 gibi) string formata çeviriyoruz
    dt_str = dogum.strftime("%d.%m.%Y")
    kt_str = kaza.strftime("%d.%m.%Y")
    ibt_str = ise_baslama.strftime("%d.%m.%Y")
    ht_str = hesap.strftime("%d.%m.%Y")

    m_gecmis = {int(r["Yıl"]): float(r["Maaş (TL)"]) for _, r in edit_df.iterrows() if pd.notna(r["Yıl"]) and pd.notna(r["Maaş (TL)"])}
    
    calc = AktueryaHesaplayici(isim, cinsiyet, dt_str, kt_str, ibt_str, ht_str, mal, kusur, m_gecmis, asgari)
    rapor = calc.hesapla()
    
    st.success("Rapor Hazır!")
    st.text_area("Önizleme", rapor, height=450)
    
    word_file = create_word_doc(rapor)
    st.download_button("📄 Raporu Word Olarak İndir", word_file, f"{isim.replace(' ', '_')}_Rapor.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")